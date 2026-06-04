"""
Parser module for extracting SEO page data from Word documents.
Supports both single-page and multi-page documents.
"""
import re
from pathlib import Path
from typing import List, Tuple, Optional
from docx import Document
from models import SEOPageData, FAQ


class SEODocumentParser:
    """Parser for SEO Word documents with multi-page support."""

    def __init__(self, doc_path: str):
        """
        Initialize parser with Word document path.

        Args:
            doc_path: Path to the .docx file
        """
        self.doc_path = Path(doc_path)
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")
        if not self.doc_path.suffix.lower() == ".docx":
            raise ValueError(f"File must be a .docx file: {doc_path}")

        self.doc = Document(str(self.doc_path))
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs]

    def _find_page_boundaries(self) -> List[int]:
        """
        Find the starting indices of each SEO page in the document.

        A new page is identified by:
        - "SEO Page X" or "SEO Page 1", "SEO Page 2", etc.
        - "Href in English:" marker

        Returns:
            List of starting indices for each page
        """
        boundaries = []
        added_indices = set()

        # Pattern 1: "SEO Page X" headers
        seo_page_pattern = re.compile(r"^SEO\s+Page\s+\d+", re.IGNORECASE)

        # Pattern 2: "Href in English:" marker
        href_pattern = re.compile(r"^Href in English\s*:", re.IGNORECASE)

        for i, para in enumerate(self.paragraphs):
            # Check if this line is a page header
            if seo_page_pattern.match(para):
                # Look for the actual "Href in English:" after this header
                # Search next 10 lines for href marker
                for j in range(i + 1, min(i + 10, len(self.paragraphs))):
                    if href_pattern.match(self.paragraphs[j]):
                        if j not in added_indices:
                            boundaries.append(j)
                            added_indices.add(j)
                        break
            # Also detect pages that start directly with "Href in English:"
            elif href_pattern.match(para):
                # Only add if not already added
                if i not in added_indices:
                    boundaries.append(i)
                    added_indices.add(i)

        return sorted(boundaries)

    def _find_field_value(self, field_name: str, start_index: int, end_index: Optional[int] = None) -> str:
        """
        Find value for a field by searching for field_name pattern within bounds.

        Args:
            field_name: Name of the field to search for (case-insensitive)
            start_index: Index to start searching from
            end_index: Index to stop searching at (exclusive), None for end of doc

        Returns:
            Extracted value or empty string if not found
        """
        pattern = re.compile(rf"^{re.escape(field_name)}\s*:\s*(.*)$", re.IGNORECASE)

        if end_index is None:
            end_index = len(self.paragraphs)

        for i in range(start_index, end_index):
            para = self.paragraphs[i]
            match = pattern.match(para)
            if match:
                return match.group(1).strip()

        return ""

    def _extract_customizable_section(self, start_index: int, end_index: Optional[int] = None) -> Tuple[str, str, str]:
        """
        Extract subtitle, title, and description from Customizable section.

        Args:
            start_index: Index to start searching from
            end_index: Index to stop searching at (exclusive)

        Returns:
            Tuple of (subtitle, title, description)
        """
        subtitle = ""
        title = ""
        description_lines = []

        if end_index is None:
            end_index = len(self.paragraphs)

        customizable_index = -1

        for i in range(start_index, end_index):
            para = self.paragraphs[i]
            if "customizable" in para.lower() and "food amigos" in para.lower():
                customizable_index = i
                break

        if customizable_index == -1:
            return subtitle, title, ""

        for i in range(customizable_index + 1, end_index):
            para = self.paragraphs[i]

            if not para:
                continue

            if para.lower().startswith("subtitle"):
                # Handle both "Subtitle: X" and "Subtitle: X Title: Y" on same line
                match = re.match(r"^subtitle\s*:\s*(.*)$", para, re.IGNORECASE)
                if match:
                    rest = match.group(1).strip()
                    # Check if Title is on the same line
                    title_match = re.match(r"^(.*?)\s+Title\s*:\s*(.*)$", rest, re.IGNORECASE)
                    if title_match:
                        subtitle = title_match.group(1).strip()
                        title = title_match.group(2).strip()
                    else:
                        subtitle = rest
                continue

            if para.lower().startswith("title:"):
                match = re.match(r"^title:\s*(.*)$", para, re.IGNORECASE)
                if match:
                    title = match.group(1).strip()
                continue

            if para.lower().startswith("description:"):
                match = re.match(r"^description:\s*(.*)$", para, re.IGNORECASE)
                if match:
                    first_line = match.group(1).strip()
                    if first_line:
                        description_lines.append(first_line)

                for j in range(i + 1, end_index):
                    next_para = self.paragraphs[j].strip()
                    if not next_para:
                        continue
                    if next_para.lower().startswith("internal links"):
                        break
                    if next_para.lower().startswith("faq"):
                        break
                    if next_para.lower().startswith("href in english"):
                        break
                    description_lines.append(next_para)
                break

        description = "\n\n".join(description_lines).strip()

        return subtitle, title, description

    def _extract_faqs(self, start_index: int, end_index: Optional[int] = None) -> List[FAQ]:
        """
        Extract FAQ items from document section.

        Args:
            start_index: Index to start searching from
            end_index: Index to stop searching at (exclusive)

        FAQ format in document:
        Title: <question>
        Description: <answer>

        Returns:
            List of FAQ objects
        """
        faqs = []

        if end_index is None:
            end_index = len(self.paragraphs)

        faq_start_index = -1

        # Try to find "Internal Links" marker first
        for i in range(start_index, end_index):
            para = self.paragraphs[i]
            if "internal links" in para.lower():
                faq_start_index = i + 1
                break

        # If no "Internal Links" found, look for "FAQ SECTION" or "FAQ" marker
        if faq_start_index == -1:
            for i in range(start_index, end_index):
                para = self.paragraphs[i]
                if "faq section" in para.lower() or (para.lower().startswith("faq") and "block" in para.lower()):
                    faq_start_index = i + 1
                    break

        if faq_start_index == -1:
            return faqs

        current_question = None

        for i in range(faq_start_index, end_index):
            para = self.paragraphs[i]

            if not para:
                continue

            # Stop if we hit the next page
            if para.lower().startswith("href in english"):
                break

            if para.lower().startswith("title"):
                # Handle both "Title: X" and "Title: X Description: Y" on same line
                match = re.match(r"^title\s*:\s*(.*)$", para, re.IGNORECASE)
                if match:
                    rest = match.group(1).strip()
                    # Check if Description is on the same line
                    desc_match = re.match(r"^(.*?)\s+Description\s*:\s*(.*)$", rest, re.IGNORECASE)
                    if desc_match:
                        # Both on same line - create FAQ immediately
                        question = desc_match.group(1).strip()
                        answer = desc_match.group(2).strip()
                        faqs.append(FAQ(question=question, answer=answer))
                        current_question = None
                    else:
                        # Only Title on this line
                        current_question = rest
                continue

            if para.lower().startswith("description"):
                if current_question:
                    match = re.match(r"^description\s*:\s*(.*)$", para, re.IGNORECASE)
                    if match:
                        answer = match.group(1).strip()
                        faqs.append(FAQ(question=current_question, answer=answer))
                        current_question = None

        return faqs

    def _parse_single_page(self, start_index: int, end_index: Optional[int] = None) -> SEOPageData:
        """
        Parse a single SEO page from the document.

        Args:
            start_index: Starting index of the page
            end_index: Ending index of the page (exclusive), None for end of doc

        Returns:
            SEOPageData object containing all extracted fields
        """
        href = self._find_field_value("Href in English", start_index, end_index)
        page_name = self._find_field_value("Name in English", start_index, end_index)
        seo_title = self._find_field_value("Title", start_index, end_index)
        seo_description = self._find_field_value("Description", start_index, end_index)

        # Ensure href starts with '/' for Food Amigo validation
        if href and not href.startswith('/'):
            href = '/' + href

        subtitle, title, description = self._extract_customizable_section(start_index, end_index)

        faqs = self._extract_faqs(start_index, end_index)

        seo_data = SEOPageData(
            href=href,
            page_name=page_name,
            seo_title=seo_title,
            seo_description=seo_description,
            subtitle=subtitle,
            title=title,
            description=description,
            faqs=faqs
        )

        return seo_data

    def parse(self) -> SEOPageData:
        """
        Parse the Word document and extract the FIRST SEO page data.

        For backward compatibility. Use parse_all() for multi-page documents.

        Returns:
            SEOPageData object containing all extracted fields

        Raises:
            ValueError: If no pages found
        """
        pages = self.parse_all()
        if not pages:
            raise ValueError("No SEO pages found in document")
        return pages[0]

    def parse_all(self) -> List[SEOPageData]:
        """
        Parse the Word document and extract ALL SEO pages.

        Returns:
            List of SEOPageData objects, one for each page in the document

        Raises:
            ValueError: If no pages found
        """
        boundaries = self._find_page_boundaries()

        if not boundaries:
            raise ValueError("No SEO pages found in document. Each page must start with 'Href in English:'")

        pages = []

        for i, start_idx in enumerate(boundaries):
            # Determine end index (start of next page or end of document)
            end_idx = boundaries[i + 1] if i + 1 < len(boundaries) else None

            try:
                page_data = self._parse_single_page(start_idx, end_idx)
                pages.append(page_data)
            except Exception as e:
                print(f"Warning: Failed to parse page {i + 1} starting at line {start_idx}: {e}")
                continue

        if not pages:
            raise ValueError("Failed to parse any valid pages from document")

        return pages


def parse_seo_document(doc_path: str) -> SEOPageData:
    """
    Convenience function to parse an SEO document (single page).

    For backward compatibility. Use parse_seo_document_all() for multi-page.

    Args:
        doc_path: Path to the .docx file

    Returns:
        SEOPageData object (first page if multiple exist)

    Raises:
        FileNotFoundError: If document doesn't exist
        ValueError: If document format is invalid
    """
    parser = SEODocumentParser(doc_path)
    return parser.parse()


def parse_seo_document_all(doc_path: str) -> List[SEOPageData]:
    """
    Parse an SEO document and extract ALL pages.

    Args:
        doc_path: Path to the .docx file

    Returns:
        List of SEOPageData objects

    Raises:
        FileNotFoundError: If document doesn't exist
        ValueError: If document format is invalid
    """
    parser = SEODocumentParser(doc_path)
    return parser.parse_all()


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parser.py <path_to_docx>")
        sys.exit(1)

    doc_path = sys.argv[1]

    try:
        pages = parse_seo_document_all(doc_path)

        print("=" * 80)
        print(f"FOUND {len(pages)} SEO PAGE(S) IN DOCUMENT")
        print("=" * 80)

        for page_num, data in enumerate(pages, 1):
            print(f"\n{'=' * 80}")
            print(f"PAGE {page_num} OF {len(pages)}")
            print("=" * 80)
            print(f"\nHref: {data.href}")
            print(f"Page Name: {data.page_name}")
            print(f"SEO Title: {data.seo_title}")
            print(f"SEO Description: {data.seo_description[:80]}...")
            print(f"\nSubtitle: {data.subtitle}")
            print(f"Title: {data.title[:80]}..." if len(data.title) > 80 else f"Title: {data.title}")
            print(f"Description: {len(data.description)} characters")
            print(f"\nFAQs: {len(data.faqs)} items")
            for i, faq in enumerate(data.faqs, 1):
                print(f"  {i}. {faq.question}")

            print("\n" + "-" * 80)
            print("VALIDATION")
            print("-" * 80)

            if data.is_valid():
                print("[OK] All required fields present")
            else:
                print("[ERROR] Missing fields:")
                for field in data.validate():
                    print(f"  - {field}")

        print("\n" + "=" * 80)
        print(f"SUMMARY: Successfully parsed {len(pages)} page(s)")
        print("=" * 80)

    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
